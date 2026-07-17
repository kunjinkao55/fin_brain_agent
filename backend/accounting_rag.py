"""
FinBrain 多知识库 RAG 模块 — 用户上传文档 → 解析 → 切片 → 向量检索

支持格式：PDF (pdfplumber), DOCX (python-docx), TXT, MD
嵌入模型：ONNX MiniLM
多知识库：每个 kb = 一个 ChromaDB collection，未来可扩（会计准则/行业研报/交易策略）

设计原则（遵循项目技术审计标准）：
- 懒初始化：首次调用时才加载模型 + 打开集合
- 线程安全：double-check locking
- 显式错误处理：解析失败/空文档/编码问题不静默
- 禁止模块级副作用：不在 import 时读写文件或下载模型
"""

import os
import json
import uuid
import logging
import threading
from datetime import datetime
from typing import Optional

import chromadb

logger = logging.getLogger(__name__)

# ---- 配置 ----
_DB_DIR = os.path.join(os.path.dirname(__file__), "..", "data", "raw", "chroma")
_UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "..", "data", "uploads")
_CHUNK_SIZE = 500       # 每块约500字
_CHUNK_OVERLAP = 50     # 块间重叠50字

# 内置知识库定义：{kb_name: {display_name, description}}
_BUILTIN_KBS = {
    "accounting": {"display_name": "会计准则", "description": "企业会计准则/IFRS/财务分析实务"},
    "industry":   {"display_name": "行业研报", "description": "行业研究/产业链分析/竞争格局"},
    "trading":    {"display_name": "交易策略", "description": "交易系统/风控/行为金融"},
}

# kb_registry 记录所有知识库的元信息（存入 ChromaDB 的 _kb_registry collection）
_REGISTRY_COLLECTION = "_kb_registry"

# ---- 懒初始化（double-check locking） ----
_client: Optional[chromadb.PersistentClient] = None
_embed_fn = None
_init_lock = threading.Lock()
_initialized = False


def _ensure_init():
    """线程安全的懒初始化。首次调用时加载嵌入模型。"""
    global _client, _embed_fn, _initialized

    if _initialized:
        return

    with _init_lock:
        if _initialized:
            return

        try:
            os.makedirs(_DB_DIR, exist_ok=True)
            os.makedirs(_UPLOAD_DIR, exist_ok=True)

            _client = chromadb.PersistentClient(path=_DB_DIR)

            try:
                _embed_fn = chromadb.utils.embedding_functions.ONNXMiniLM_L6_V2()
            except Exception:
                _embed_fn = chromadb.utils.embedding_functions.DefaultEmbeddingFunction()

            # 注册内置知识库（幂等）
            _sync_registry()

            _initialized = True
            logger.info("Multi-KB RAG initialized: embedder=%s", type(_embed_fn).__name__)
        except Exception as e:
            logger.exception("Multi-KB RAG initialization failed: %s", e)
            raise RuntimeError(f"多知识库RAG初始化失败: {e}") from e


def _get_kb_collection(kb_name: str):
    """获取指定知识库的 ChromaDB collection（自动创建）。"""
    _ensure_init()
    col_name = f"kb_{kb_name}"
    return _client.get_or_create_collection(
        name=col_name,
        embedding_function=_embed_fn,
        metadata={"kb_name": kb_name},
    )


def _sync_registry():
    """将内置知识库注册到 _kb_registry（幂等）。"""
    try:
        reg = _client.get_or_create_collection(
            name=_REGISTRY_COLLECTION,
            metadata={"description": "知识库注册表"},
        )
    except Exception:
        return  # registry 创建失败不影响使用

    existing = reg.get()
    existing_ids = set(existing.get("ids", []))

    for kb_name, info in _BUILTIN_KBS.items():
        if kb_name not in existing_ids:
            try:
                reg.add(
                    documents=[info["description"]],
                    ids=[kb_name],
                    metadatas=[{
                        "name": kb_name,
                        "display_name": info["display_name"],
                        "description": info["description"],
                        "created": datetime.now().strftime("%Y-%m-%d %H:%M"),
                        "builtin": True,
                    }],
                )
            except Exception:
                pass  # 并发创建时可能冲突，忽略


# ============================================================
#  文档解析（不变）
# ============================================================

def _parse_pdf(file_path: str) -> str:
    """用 pdfplumber 提取 PDF 文本。失败时回退到 pypdf。"""
    text_parts = []
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        if text_parts:
            return "\n\n".join(text_parts)
    except Exception as e:
        logger.warning("pdfplumber extraction failed for %s: %s, falling back to pypdf",
                       os.path.basename(file_path), e)

    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n\n".join(text_parts)
    except Exception as e:
        raise RuntimeError(f"PDF解析失败 ({os.path.basename(file_path)}): {e}") from e


def _parse_docx(file_path: str) -> str:
    """用 python-docx 提取 DOCX 文本（含表格）。"""
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    text_parts.append(" | ".join(cells))

        return "\n\n".join(text_parts)
    except Exception as e:
        raise RuntimeError(f"DOCX解析失败 ({os.path.basename(file_path)}): {e}") from e


def _parse_txt(file_path: str) -> str:
    """读取纯文本文件，尝试多种编码。"""
    for encoding in ["utf-8", "gbk", "gb2312", "gb18030", "latin-1"]:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise RuntimeError(f"无法识别文件编码: {os.path.basename(file_path)}")


def _parse_file(file_path: str, filename: str) -> str:
    """根据扩展名分发到对应解析器。返回纯文本。"""
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        text = _parse_pdf(file_path)
    elif ext in (".docx", ".doc"):
        text = _parse_docx(file_path)
    elif ext in (".txt", ".md", ".csv", ".json"):
        text = _parse_txt(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}（支持 PDF/DOCX/TXT/MD/CSV/JSON）")

    if not text or not text.strip():
        raise ValueError(f"文件内容为空或无法提取文本: {filename}")

    return text.strip()


# ============================================================
#  文本切片（不变）
# ============================================================

def _split_long_paragraph(text: str, chunk_size: int, overlap: int) -> list[str]:
    """将超长段落按句子边界切分，尽量保持语义完整。"""
    sentences = []
    current = ""
    for char in text:
        current += char
        if char in "。！？；\n":
            sentences.append(current)
            current = ""
    if current.strip():
        sentences.append(current)

    chunks = []
    current_chunk = ""
    for sent in sentences:
        if len(current_chunk) + len(sent) <= chunk_size:
            current_chunk += sent
        else:
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            if len(sent) > chunk_size:
                for i in range(0, len(sent), chunk_size - overlap):
                    chunks.append(sent[i:i + chunk_size])
                current_chunk = ""
            else:
                current_chunk = sent
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    return chunks


def _chunk_text(text: str, chunk_size: int = _CHUNK_SIZE,
                overlap: int = _CHUNK_OVERLAP) -> list[str]:
    """将文本按段落→句子→字符三级切分为重叠块。"""
    paragraphs = text.split("\n\n")
    chunks = []
    current = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        if len(current) + len(para) <= chunk_size:
            current += para + "\n\n"
        else:
            if current.strip():
                chunks.append(current.strip()[:chunk_size + overlap])
            if len(para) > chunk_size:
                sub_chunks = _split_long_paragraph(para, chunk_size, overlap)
                chunks.extend(sub_chunks)
                current = ""
            else:
                current = para + "\n\n"

    if current.strip():
        chunks.append(current.strip()[:chunk_size + overlap])

    return [c for c in chunks if len(c) >= 20]


# ============================================================
#  预置会计准则知识
# ============================================================

_ACCOUNTING_SEED = [
    # ---- 收入与利润 ----
    {
        "id": "seed_revenue_recognition",
        "title": "收入确认五步法",
        "content": (
            "IFRS 15 / CAS 14 收入确认五步法：(1)识别与客户的合同——合同须具有商业实质、"
            "双方批准且承诺履行义务、付款条款明确；(2)识别合同中的履约义务——可明确区分的"
            "商品或服务承诺，判断标准：客户可单独获益+承诺在合同中可单独区分；"
            "(3)确定交易价格——包括可变对价(折扣/返利)、重大融资成分、非现金对价；"
            "(4)将交易价格分摊至各履约义务——按单独售价比例分配，无法直接观察时用"
            "预计成本加毛利法或余值法；(5)在满足履约义务时确认收入——时点义务以控制权转移为标志"
            "(验收/签收/风险转移)，时段义务按完工进度(产出法/投入法)。"
            "财报分析要点：提前确认收入的手段包括经销商压货(渠道填塞)、捆绑销售分拆不当、"
            "将代理收入按总额法而非净额法确认。警惕应收增速远超营收增速(表明赊销驱动增长)。"
        ),
    },
    {
        "id": "seed_asset_impairment",
        "title": "资产减值测试",
        "content": (
            "CAS 8 资产减值：企业应在资产负债表日判断资产是否存在减值迹象。减值迹象包括："
            "市价大幅下跌、经营环境重大不利变化、利率上升导致折现率升高、资产陈旧过时、"
            "实际业绩低于预期。测试方法：比较账面价值与可收回金额(公允价值-处置费用 vs 使用价值，取孰高)。"
            "使用价值=预计未来现金流折现(DCF)，折现率应反映当前市场对货币时间价值和资产特定风险的评价。"
            "关键：资产组(CGU)划分——独立产生现金流入的最小资产组合；商誉必须每年测试不论是否有减值迹象；"
            "总部资产(研发中心/管理楼)按合理基础分摊至CGU。财报分析要点：减值损失大幅波动是盈余管理的"
            "常见信号——好年景多计提做低基数(洗大澡)，差年景少计提保利润。关注商誉占比高的公司，"
            "商誉减值往往比预期晚1-2年(管理层乐观偏差)。"
        ),
    },
    {
        "id": "seed_goodwill",
        "title": "商誉减值与并购",
        "content": (
            "商誉产生于非同一控制下企业合并，购买方支付的对价超过被购买方可辨认净资产公允价值份额的部分。"
            "CAS 20 要求每年年终对商誉进行减值测试，不得摊销。实务中，商誉减值是最常见的爆雷点："
            "并购时高估(用收益法而非市场法估值，增长率假设乐观)，3-5年后业绩不达预期时巨额减值。"
            "财报分析要点：(1)商誉/净资产>30%的并购驱动型公司需警惕；"
            "(2)业绩承诺期(通常3年)刚过就减值=并购时可能虚增商誉；"
            "(3)商誉减值测试的折现率(WACC)如逐年降低，可能是在回避减值——折现率越低，可收回金额越高；"
            "(4)分步收购：先买少量股权(不形成控制)，估值低→再增持取得控制权时形成大额商誉=操纵。"
        ),
    },
    # ---- 资产与负债 ----
    {
        "id": "seed_financial_instruments",
        "title": "金融工具分类与计量",
        "content": (
            "CAS 22 将金融资产分为三类：(1)摊余成本(AC)——业务模式为收取合同现金流+现金流仅为本金和利息"
            "(SPPI测试通过)，如普通债券/应收款；(2)以公允价值计量且其变动计入其他综合收益(FVOCI)——"
            "业务模式既收取合同现金流又出售，如可转债/部分权益投资指定；(3)以公允价值计量且其变动计入"
            "当期损益(FVTPL)——交易性/衍生工具/不满足SPPI的。金融负债分为摊余成本和FVTPL。"
            "财报分析要点：(1)FVTPL资产占比高→利润波动大(公允价值变动直接进利润表)；"
            "(2)FVOCI权益投资：公允价值变动不进利润表但卖出的累计OCI要转留存收益(不转利润)，"
            "意味着卖出亏损不会体现在净利润中；(3)应收款项融资(如票据贴现)终止确认条件："
            "转移了几乎所有的风险和报酬(追索权是关键)。"
        ),
    },
    {
        "id": "seed_lease",
        "title": "租赁负债与使用权资产",
        "content": (
            "CAS 21(2018修订，与IFRS 16趋同)：承租人不再区分经营租赁和融资租赁，几乎所有租赁"
            "均需确认使用权资产和租赁负债(豁免：短期租赁≤12个月+低价值资产租赁)。"
            "使用权资产=租赁负债+已付预付+初始直接费用-租赁激励；租赁负债=租赁付款额折现"
            "(折现率用增量借款利率(承租人)或内含利率(出租人))。租赁付款额包括：固定付款+"
            "可变租赁付款额(取决于指数或利率的部分)+合理确定的购买选择权+终止租赁罚款。"
            "财报分析要点：(1)重资产零售/航空行业租赁负债巨大，用净债务/EBITDA(含租赁)替代"
            "传统资产负债率判断真实杠杆；(2)租赁负债的折现率如异常偏高→压低负债额→美化报表；"
            "(3)关联方租赁定价低于市场价→相当于大股东变相资金占用。"
        ),
    },
    {
        "id": "seed_inventory",
        "title": "存货计价与跌价准备",
        "content": (
            "CAS 1 存货按成本与可变现净值孰低计量。发出成本计价方法：先进先出(FIFO)、加权平均法"
            "(已取消后进先出LIFO)。存货成本包括采购成本、加工成本(料工费)和使存货达到当前位置和状态"
            "的其他成本。可变现净值=估计售价-至完工估计成本-估计销售费用。"
            "财报分析要点：(1)毛利率趋势与存货计价方法联动分析——FIFO下存货余额接近当前市价(资产负债表准)，"
            "但销货成本低→毛利率偏高→利润虚高(通胀环境)；(2)存货跌价准备是盈余管理重灾区，"
            "存货增加远超营收增加→可能是滞销(需计提跌价)或提前备货(看下季订单)；(3)存货周转天数"
            "延长+毛利率下降=产品竞争力下降的经典组合信号；(4)农业/养殖企业存货盘点困难(审计受限)"
            "是造假高发领域。"
        ),
    },
    # ---- 现金流量 ----
    {
        "id": "seed_cashflow_quality",
        "title": "经营现金流质量分析",
        "content": (
            "现金流质量是判断财报真实性的核心维度。经营现金流/净利润(OFC/NI)是关键指标："
            "长期低于0.7说明利润含金量不足，低于0.5是危险信号。分析框架：(1)经营性应收项目增加"
            "=应收账款/应收票据增长>收入增长→激进赊销；预付账款大幅增长→可能关联方资金占用；"
            "(2)经营性应付项目增加=应付账款增长远高于营业成本增长→可能拖欠供应商回笼现金(看应付周转天数)，"
            "不可持续；(3)存货增加占用现金(存货/营收比攀升=滞销)；(4)折旧摊销＞购建固定资产"
            "→长期资产更新不足=短期化行为。关注经营性现金流与自由现金流的差距："
            "自由现金流=经营现金流-资本支出(购建固定资产/无形资产支付的现金)。"
            "FCF持续为正且增长的公司有真正的股东回报能力(分红/回购)。"
        ),
    },
    {
        "id": "seed_rd_capitalization",
        "title": "研发支出资本化",
        "content": (
            "CAS 6 无形资产：研究阶段支出全部费用化；开发阶段支出同时满足5个条件时可资本化——"
            "(1)完成该无形资产以使其能够使用或出售在技术上具有可行性；(2)具有完成并使用或出售的意图；"
            "(3)无形资产产生经济利益的方式(自身使用/出售)可证明；(4)有足够的技术/财务资源支持开发完成；"
            "(5)开发阶段的支出能够可靠计量。实务中资本化条件是分水岭：取得软件著作权/专利证书是"
            "标志性节点。财报分析要点：(1)研发资本化率(资本化/总研发支出)在同行业内横向比较，"
            "异常偏高的公司可能通过资本化美化利润——资本化意味着研发支出不进利润表，而是形成无形资产"
            "分期摊销；(2)资本化率的变动趋势比绝对值更重要：突然大幅提高→可能在保利润；"
            "(3)计算机/医药/汽车行业研发资本化率差异大，需在同行业同生命周期阶段比较。"
        ),
    },
    {
        "id": "seed_consolidation",
        "title": "合并报表范围与控制权",
        "content": (
            "CAS 33 合并财务报表：控制=权力+可变回报+权力影响回报的能力。控制三要素：(1)对被投资方的权力"
            "(表决权>50%或协议/章程赋予主导权)；(2)因参与被投资方而享有可变回报(股利/利息/服务费/"
            "规模效益/品牌协同)；(3)有能力运用权力影响回报(实际主导，非被动)。"
            "财报分析要点：(1)出表/并表的边界操纵——通过结构化主体(SPV/资管计划/信托)隐藏负债，"
            "将亏损子公司通过协议安排(Variable Interest Entities)排除在合并范围外；"
            "(2)少数股东损益占比异常高→说明母公司实际只控制但只有少量股权→可能通过少数股权"
            "腾挪利润；(3)合并范围变化(收购/处置子公司)导致的收入增长不是有机增长，细分可比口径。"
        ),
    },
    # ---- 估值与财务分析 ----
    {
        "id": "seed_roe_dupont",
        "title": "ROE 杜邦三因子分解",
        "content": (
            "ROE = 净利润/净资产 = (净利润/营收) × (营收/总资产) × (总资产/净资产) = 净利率 × 资产周转率 × 权益乘数。"
            "三因子含义：(1)净利率=盈利能力——取决于毛利率(产品竞争力)和费用管控(管理效率)；"
            "毛利率>40%为优秀，20-40%为正常，<20%需警惕(除非是高周转模式)；"
            "(2)资产周转率=运营效率——取决于产能利用率、存货管理、应收账款回收速度；"
            "轻资产模式(互联网/软件)周转率天然高，重资产(制造/电力)天然低，不可跨行业对比；"
            "(3)权益乘数=财务杠杆——适度负债可放大ROE，但乘数>3时风险显著上升"
            "(即资产负债率>67%，每1元净资产撬动>3元资产)。"
            "哪种ROE更可持续？高净利率驱动的ROE(茅台/海天)最可持续(品牌壁垒)，"
            "高杠杆驱动的ROE(地产/银行)具有周期性，高周转驱动的ROE(零售/电商)需持续运营投入。"
            "ROE>15%连续3年且三因子结构稳定=优质公司。"
        ),
    },
    {
        "id": "seed_related_party",
        "title": "关联交易与资金占用",
        "content": (
            "关联方交易是财务造假的高发地带。CAS 36 定义了关联方：控制/共同控制/重大影响的关系，"
            "包括母子公司、同一控制下的兄弟公司、关键管理人员及关系密切家庭成员控制的企业。"
            "重点关注的关联交易类型：(1)商品购销——关联采购额占比过高→可能转移定价(高价采购=利润输送)"
            "或供应商集中风险；(2)资金拆借——其他应收款中关联方余额异常增大→实质是大股东资金占用；"
            "(3)担保——为关联方提供巨额担保(关注或有负债附注)；(4)资产收购/出售——向关联方高溢价"
            "收购资产→套现上市公司资金；(5)共同投资——与关联方共同出资设立基金/公司→隐藏投资风险。"
            "分析要点：关联方应收应付余额>当年关联交易额的50%→可能存在未结算的非经营性资金往来。"
            "同时关注前五大客户/供应商是否为隐形关联方(同名/同注册地/同电话)。"
        ),
    },
    {
        "id": "seed_nonrecurring",
        "title": "非经常性损益与扣非净利润",
        "content": (
            "非经常性损益指与正常经营业务无直接关系、或虽相关但性质和发生频率影响报表使用者"
            "合理评价企业盈利能力的交易或事项。常见项目：(1)非流动资产处置损益(卖房/卖地/卖子公司)"
            "→一次性的资产处置收益；(2)政府补助(与经营活动无关的)；(3)债务重组收益(债权人让步)"
            "→公司主业亏损但靠减免债务扭亏；(4)公允价值变动损益(投资性房地产/交易性金融资产)"
            "→市场波动导致的持有利得/损失；(5)单独进行减值测试的应收款项减值准备转回。"
            "财报分析要点：用扣非净利润而非净利润判断持续经营能力。扣非/净利润<0.7→利润质量差"
            "(超30%利润来自非经常项目)。连续多年扣非净利润为负但净利润为正→典型的粉饰报表，"
            "这类公司主营业务已失去竞争力。扣非口径下连续3年亏损≈主业实质上已经ST。"
        ),
    },
    {
        "id": "seed_deferred_tax",
        "title": "递延所得税资产与负债",
        "content": (
            "递延所得税产生于会计准则与税法对收入/费用的确认时点差异(暂时性差异)。"
            "递延所得税资产(DTA)：会计上已确认费用但税前不能抵扣(如资产减值准备/预提费用/"
            "可抵扣亏损)，未来可抵税=企业的一项资产。递延所得税负债(DTL)：会计上确认了收入但"
            "税前可递延纳税(如固定资产加速折旧(税法)>会计折旧=多抵税，以后要补)，"
            "未来要多交税=一项负债。财报分析要点：(1)DTA/净资产>10%→依赖未来盈利实现抵税，"
            "如果公司持续亏损则DTA需计提减值(估值)；(2)DTA急剧增长→可能在用可抵扣亏损做大资产"
            "(虚增净资产)；(3)高DTL→意味着过去加速折旧等税收优惠很大，未来税负会加重"
            "(现金流预警)；(4)会计利润远大于应纳税所得额→永久性差异(免税收入/不可抵扣支出)"
            "过大，有效税率异常低需要解释(税收优惠/利润转移/造假)。"
        ),
    },
    {
        "id": "seed_earnings_management",
        "title": "盈余管理的常见手法",
        "content": (
            "盈余管理(合法但扭曲报表)与财务造假(违法)的边界有时模糊，分析时需保持警惕。"
            "常见手法：(1)收入端——渠道填塞(季末压货/放宽信用条件)、提前确认(未达验收条件即确认)、"
            "总额法vs净额法选择(代理业务按总额法虚增收入)、虚构返利(冲减收入)；"
            "(2)费用端——费用资本化(研发/借款费用/长期待摊)、少提折旧(延长折旧年限)、"
            "少提坏账(放宽信用政策但准备率不变)、削减研发/广告等酌量性支出(牺牲长期竞争力保短期利润)；"
            "(3)资产负债表端——表外负债(经营租赁vs融资租赁/结构化主体)、"
            "商誉减值的时机选择(差年景一次性大额减值=洗澡，好年景少提=平滑)、"
            "售后回租(利用资产评估增值做高净资产)；(4)现金流——将经营流出分类到投资/筹资活动"
            "(如将经营租赁付款归入投资活动)、应收票据贴现不终止确认(实际上有追索权)。"
            "分析框架：应收账款/营收增速差、存货/营收增速差、经营现金流/净利润比值三指标联合看，"
            "连续两个季度恶化就是预警信号。"
        ),
    },
    {
        "id": "seed_industry_metrics",
        "title": "行业特定财务指标",
        "content": (
            "不同行业的会计要点和关键指标不同：(1)银行——关注净息差(NIM)、不良贷款率、拨备覆盖率、"
            "核心一级资本充足率；拨备是银行最大的盈余调节项(拨备少提→利润高→未来补提压力大)。"
            "(2)地产——关注合同负债(预售房款)、土地增值税拨备充足性、利息资本化率、"
            "剔除合同负债后的净负债率；合作项目(联合营)的表外负债风险。"
            "(3)制造——关注产能利用率、固定资产成新率、折旧方法变更、存货跌价准备充分度。"
            "(4)医药——研发管线(研发支出资本化分阶段)、销售费用(CSO模式下合规风险)、"
            "商誉(并购驱动的增长是否可持续)。(5)消费——关注渠道库存(可通过应收账款/收入比间接观察)、"
            "品牌投入(广告费/营收)、经销商体系稳定性(前五大客户集中度高可能是直销占比提升)。"
            "(6)科技——关注研发投入强度(研发/营收)、人均产出、股权激励费用(非现金支出但摊薄每股收益)。"
            "跨行业分析时，始终将公司与同行业可比公司的指标做横向对比，而非绝对阈值判断。"
        ),
    },
]


# ============================================================
#  预置知识库种子函数
# ============================================================

def seed_accounting_kb() -> dict:
    """将会计准则预置知识写入 accounting 知识库（幂等：已有则不重复添加）。

    Returns:
        {"seeded": N, "skipped": M} — 新增条数和跳过条数
    """
    _ensure_init()
    col = _get_kb_collection("accounting")

    # 检查已有数据
    existing = col.get()
    existing_ids = set(existing.get("ids", []))

    new_docs = []
    new_ids = []
    new_metas = []
    skipped = 0

    for item in _ACCOUNTING_SEED:
        if item["id"] in existing_ids:
            skipped += 1
            continue
        new_ids.append(item["id"])
        new_docs.append(f"【{item['title']}】{item['content']}")
        new_metas.append({
            "doc_id": item["id"],
            "filename": f"[预置]{item['title']}",
            "chunk_index": 0,
            "total_chunks": 1,
            "upload_date": "2026-07-15 (built-in)",
            "char_count": len(item["content"]),
            "source": "builtin",
        })

    if new_docs:
        col.add(documents=new_docs, ids=new_ids, metadatas=new_metas)
        logger.info("Seeded %d accounting KB entries (skipped %d existing)", len(new_docs), skipped)

    return {"seeded": len(new_docs), "skipped": skipped}


# ---- 行业分析模板 ----

_INDUSTRY_TEMPLATES = [
    {"id": "ind_semiconductor", "title": "半导体行业模板",
     "content": (
        "【行业周期】半导体是强周期行业，完整周期约3-4年。周期驱动:下游需求(AI/手机/PC)+产能投放+库存周期。"
        "2024-2026处于AI驱动上行周期，但传统消费电子仍疲弱。周期位置判断比当期利润更重要——利润峰值往往是估值低点(PE<15)，利润谷底往往是估值高点(PE>50)。"
        "【估值规律】设计/设备:PE 30-60倍(成长溢价)；制造/封测:PE 15-35倍(重资产周期)。A股半导体历史PE中枢约40-50倍。"
        "封测(长电/通富/华天):PE 15-25倍成熟期合理，先进封装占比高可给30-40倍。日月光(全球第一)PE约20倍，安靠PE约25倍。"
        "周期底部:PE虚高→用PB或EV/EBITDA更合理。周期顶部:PE看起来便宜→这是陷阱，利润即将下行。"
        "【竞争格局】全球:台积电>三星>英特尔。国内:中芯(制造)/韦尔(设计)/长电(封测)。国产替代长期逻辑但技术差距仍需5-10年。"
        "先进封装(Chiplet/HBM):台积电主导，长电/通富国内领先。壁垒:百亿级资本开支+3-5年客户认证+技术专利。"
        "【产业链位置】上游:设备(北方华创)+材料(沪硅);中游:制造(中芯)+封测(长电);下游:芯片设计(韦尔/兆易)。封测毛利率天然偏低(15-25%)。"
        "【关键指标】全球半导体销售额(WSTS月度)、晶圆厂产能利用率、库存周转天数、先进封装收入占比、北美设备BB率。"
    )},
    {"id": "ind_power", "title": "电力/能源行业模板",
     "content": (
        "【行业周期】火电:强周期，利润与煤价反向。煤价↑→利润↓，煤价↓→利润↑，典型周期2-3年。水电:弱周期，来水量决定利润。新能源:成长型，渗透率提升驱动。电力需求偏刚性(GDP增速×1.2倍)。"
        "【估值规律】火电:PE 10-18倍(周期股，PE最低=利润高点=周期顶部!)。水电:PE 15-25倍(稳定现金流溢价)。新能源:PE 20-35倍(成长溢价)。PB中枢1.5-2.5倍。股息率3-5%正常。"
        "火电PE反向指标:PE<10倍→利润见顶风险，PE>30倍→利润见底机会。"
        "【竞争格局】五大发电集团寡头竞争，市占率约50%。新进入者:新能源民企+央企背景地方能源。火电壁垒:审批+电网接入+资本规模(单机组数十亿)。"
        "【产业链位置】上游:煤炭+天然气;中游:发电企业;下游:电网+售电。发电企业夹在燃料和电价之间，利润弹性来自煤价。"
        "【关键指标】利用小时数(>5000优)、标煤价(秦港5500大卡)、市场化电价比例、风光装机增速、资产负债率(>70%重资产正常)。"
    )},
    {"id": "ind_pharma", "title": "医药行业模板",
     "content": (
        "【行业周期】医药弱周期(需求刚性)，但受政策(集采/医保)和技术(新药)周期影响。创新药:管线驱动。仿制药:集采降价周期1-2年。CXO:融资周期驱动，滞后创新药融资2-3年。"
        "【估值规律】成熟药企(恒瑞/丽珠):PE 20-40倍。创新药:PS+rNPV估值(无利润)。CXO:PEG估值，增速20-30%→PE 30-50倍。中药(片仔癀):PE 25-45倍(品牌溢价)。A股医药PE中枢30-40倍。"
        "【竞争格局】创新药:国际巨头主导，国内恒瑞/百济追赶。CXO:中国成本+工程师红利。中药:政策+品牌壁垒。仿制药:集采后集中度提升。"
        "【产业链位置】上游:原料药(华海/普洛);中游:制剂+CXO;下游:医院/药店/医保。价值链:研发(80%+)>生产(30-60%)>流通(5-15%)。"
        "【关键指标】研发费用率(>10%创新驱动)、核心品种销售趋势、集采中标、在研管线阶段、一致性评价数。"
    )},
    {"id": "ind_consumer", "title": "消费品行业模板",
     "content": (
        "【行业周期】食品饮料:弱周期(必选消费)。白酒:库存周期3-4年，与商务活动相关。家电:地产+出口周期驱动。"
        "【估值规律】白酒(茅台/五粮液):PE 20-40倍(品牌溢价+高ROE)。食品(海天/伊利):PE 25-40倍。家电(美的/格力):PE 10-18倍(成熟制造)。A股消费PE中枢25-35倍。高ROE(>20%)+低负债(<30%)=优质消费品。"
        "【竞争格局】白酒:茅台独大(高端)，五粮液/泸州老窖/洋河竞争次高端。调味品:海天领先。乳制品:伊利/蒙牛双寡头。家电:美的/格力/海尔三巨头。"
        "【产业链位置】上游:原料(粮食/原奶);中游:品牌商;下游:经销商/电商。品牌商掌握定价权(高毛利)，经销商利润薄。"
        "【关键指标】产品结构(高端占比)、渠道库存周转天数、经销商数量变化、终端动销率、复购率、广告费/营收比。"
    )},
    {"id": "ind_communication", "title": "通信/光模块行业模板",
     "content": (
        "【行业周期】光模块:AI驱动高速增长(800G/1.6T)，历史上强周期性(每代放量→饱和→杀价→新一代)。当前800G放量+1.6T导入期，景气度极高但持续性需跟踪AI资本开支。"
        "【估值规律】光模块(旭创/新易盛):PE 30-80倍(增速决定PE)，PEG<1合理。通信设备(中兴):PE 15-25倍。A股通信PE中枢25-40倍。"
        "PE与增速强相关:增速>50%→PE 50-80;增速20-50%→PE 30-50;增速<20%→PE 20-30。用单季暴增算PEG严重低估!应用2-3年复合增速。"
        "【竞争格局】全球:旭创(前三)/新易盛/天孚/Coherent/Lumentum。壁垒:技术迭代+客户认证(3-5年)+量产良率+规模成本。DAC/LPO/CPO等技术演进中，技术路线替代是核心风险。"
        "【产业链位置】上游:光芯片(源杰)+电芯片;中游:光模块组装(旭创/新易盛);下游:数据中心(Google/Meta/英伟达)。光模块毛利率20-45%。"
        "【关键指标】AI资本开支、1.6T放量节奏、毛利率(>40%优)、客户集中度、硅光/CPO进展。"
    )},
    {"id": "ind_manufacturing", "title": "制造业/新能源行业模板",
     "content": (
        "【行业周期】制造受产能周期(扩产→过剩→出清)+库存周期双驱动。新能源(锂电/光伏):产能过剩→出清阶段，毛利率承压。汽车:电动化渗透率提升中，智能化新增长点。"
        "【估值规律】成长制造(宁德/比亚迪):PEG估值，PE 20-40倍。周期制造(宝钢):PE 10-18倍+PB 0.8-1.5倍。新能源(隆基/通威):PE 10-20倍(产能过剩压缩估值)。PB比PE更稳定(重资产)。"
        "【竞争格局】锂电:宁德(全球第一35%)+比亚迪双寡头。光伏:隆基/通威/晶科产能过剩价格战。汽车:比亚迪+特斯拉+新势力。壁垒:规模效应+客户认证+技术专利+百亿级资本。"
        "【产业链位置】上游:原材料(锂/钴/硅料);中游:零部件/电芯/组件;下游:整车/电站。制造环节利润最薄(10-20%)，上游资源弹性最大。"
        "【关键指标】产能利用率(>80%优)、毛利率趋势、CAPEX/折旧比、存货周转天数、订单/合同负债增速。"
    )},
    {"id": "ind_financial", "title": "金融业分析模板",
     "content": (
        "【行业周期】银行:信用周期(不良率+拨备)+利率周期(息差)。保险:保费周期+投资收益率周期。券商:市场周期(Beta强，牛市暴增熊市暴跌)。"
        "【估值规律】银行:PB 0.5-1.2倍+ROE 8-15%。优质银行:ROE>12%+不良率<1.5%+拨备>200%→PB>1倍。保险:PEV 0.5-1.5倍。券商:PE 15-30倍，PB 1-2倍更稳定。A股银行PB中枢0.8-1.0，券商1.5-2.0。"
        "【竞争格局】银行:六大国有+股份制(招商/兴业)+城商行(宁波)，集中度高。券商:中信/华泰/国泰君安领先。保险:国寿/平安/太保三巨头。"
        "【产业链位置】金融处于经济价值链顶端，利润源于实体经济融资和投资需求。利率↓→银行息差压缩；利率↑→保险改善。"
        "【关键指标】银行:净息差(NIM)、不良率+拨备率、核心资本充足率、非利息收入比。保险:NBV、综合成本率、投资收益率。券商:两融余额、股基交易量、IPO承销额。"
    )},
    {"id": "ind_realestate", "title": "地产/建筑行业模板",
     "content": (
        "【行业周期】地产:政策周期+信贷周期+库存周期(3-5年)。当前去杠杆底部，政策放松但需求修复慢。建筑:基建投资周期+地产新开工周期。"
        "【估值规律】地产:PB 0.3-0.8倍+PE 5-15倍。PB<0.5=市场认为净资产有水分。建筑:PE 5-12倍+PB 0.8-1.5倍。A股地产PB中枢1.0-1.5(当前远低)。"
        "【竞争格局】地产:国企(保利/招商)+民企(龙湖)+城投。集中度快速提升(出险民企退出→国企份额上升)。建筑:中建/中铁/中铁建央企三巨头。"
        "【产业链位置】上游:土地+建材;中游:开发商/建筑商;下游:购房者/商业租户。开发商利润受地价/房价/融资成本三重挤压。"
        "【关键指标】地产:合同销售额+回款率、土储去化周期、综合融资成本、净负债率。建筑:新签订单增速、在手订单/收入比、PPP回款。"
    )},
]


# ---- 交易策略知识库 ----

_TRADING_TEMPLATES = [
    {
        "id": "trade_value_investing",
        "title": "价值投资策略",
        "content": (
            "【核心理念】以低于内在价值的价格买入优质公司，长期持有等待价值回归。"
            "【选股标准】ROE>15%连续3年、经营现金流/净利润>0.8、资产负债率<50%、PE<行业均值。"
            "【买入信号】PE处于5年历史分位<30% 且 基本面未恶化（ROE未趋势性下降）。"
            "【卖出信号】PE超过历史均值+1个标准差 或 基本面恶化（ROE连续2年下降>20%）。"
            "【典型标的】茅台、招商银行、美的集团、长江电力。"
            "【适用市场】熊市、震荡市。不适合：牛市末期、概念炒作期。"
        ),
    },
    {
        "id": "trade_growth_investing",
        "title": "成长股投资策略",
        "content": (
            "【核心理念】投资于高增长行业中具有竞争优势的公司，以合理价格买入成长。"
            "【选股标准】营收增速>20%连续2年、扣非净利润增速>营收增速、行业空间>千亿、市占率提升趋势。"
            "【估值方法】PEG<1.5为合理（PE/盈利增速）。PE很高不一定贵——如果增速持续，高PE会自然消化。"
            "【买入信号】PEG<1 且 季报增速未放缓（QoQ增速>0）。"
            "【卖出信号】连续2个季度营收增速<15% 或 PEG>2.5。"
            "【典型标的】宁德时代、比亚迪、中际旭创、新易盛。"
            "【致命错误】在增速放缓时还按成长股估值（戴维斯双杀）；用PE绝对值判断贵贱。"
        ),
    },
    {
        "id": "trade_cyclical",
        "title": "周期股投资策略",
        "content": (
            "【核心理念】在行业周期底部买入，在周期顶部卖出。利润最差时最值得买，利润最好时最该卖。"
            "【周期位置判断】产能利用率<70%→底部；库存见顶回落→拐点；新增产能投放→顶部风险。"
            "【买入信号】PE极高或亏损+行业龙头开始减产+库存连续3个月下降。"
            "【卖出信号】PE极低（<10倍）+新产能大量投放+媒体广泛报道行业景气。"
            "【典型标的】钢铁（宝钢）、化工（万华）、航运（中远海控）、猪肉（牧原）。"
            "【致命错误】PE低时买入（周期顶部的PE陷阱）；把周期反弹当成长；忽视产能周期（2-3年）。"
        ),
    },
    {
        "id": "trade_turnaround",
        "title": "困境反转策略",
        "content": (
            "【核心理念】投资于暂时陷入困境但核心竞争力未受损的公司，等待业绩修复。"
            "【反转信号】管理层更换+战略收缩（剥离亏损业务）+成本削减计划+核心产品销量止跌。"
            "【买入信号】扣非净利润降幅收窄（连续2季）+毛利率企稳+经营现金流转正。"
            "【卖出信号】反转失败：3个季度利润未改善 或 核心竞争力被证伪（市占率持续下降）。"
            "【仓位管理】单只不超过总仓位5%，分批建仓（第一次建仓30%+确认信号70%）。"
            "【致命错误】把价值陷阱当困境反转（市占率持续下降的不是困境，是衰落）；一次性重仓。"
        ),
    },
    {
        "id": "trade_risk_management",
        "title": "风险管理与仓位策略",
        "content": (
            "【仓位原则】单只股票≤10%（核心持仓）或≤5%（卫星仓位）。单一行业≤30%。现金≥10%。"
            "【止损纪律】基本面止损：投资逻辑被证伪→无条件卖出。估值止损：PE超过合理区间+1个标准差→减半仓。"
            "【分批建仓】底仓30%→确认信号加30%→趋势确认加40%。不一次性满仓。"
            "【卖出优先级】1.逻辑证伪(无条件) 2.估值过高(分批减) 3.发现更好标的(置换) 4.需要现金(最后)。"
            "【黑天鹅应对】单日跌幅>7%→检查基本面是否变化；跌幅>15%+基本面恶化→清仓；跌幅>15%+基本面正常→加仓。"
        ),
    },
    {
        "id": "trade_market_environment",
        "title": "市场环境与风格切换",
        "content": (
            "【牛市】重成长股+重仓位(>80%)+放松估值约束。特征：成交量放大、融资余额上升、新开户数增加。"
            "【熊市】重价值股+重防御(高股息/必选消费)+现金>30%。特征：成交量萎缩、破净股增多、IPO暂停。"
            "【震荡市】均衡配置+波段操作+PE/PB做区间交易。特征：指数在±10%内波动、无明显趋势。"
            "【风格判断】利率下行→成长股受益；利率上行→价值股受益。通胀上升→周期股+资源股。政策宽松→券商+地产。"
            "【情绪指标】融资余额/成交量比值（>10%偏热）、新基金发行量（冰点=底部信号）、沪深300波动率（>30=恐慌底）。"
        ),
    },
]

def seed_trading_kb() -> dict:
    """将交易策略模板写入 trading 知识库（幂等）。"""
    _ensure_init()
    col = _get_kb_collection("trading")
    existing = col.get()
    existing_ids = set(existing.get("ids", []))

    new_docs, new_ids, new_metas = [], [], []
    skipped = 0
    for item in _TRADING_TEMPLATES:
        if item["id"] in existing_ids:
            skipped += 1
            continue
        new_ids.append(item["id"])
        new_docs.append(f"【{item['title']}】{item['content']}")
        new_metas.append({
            "doc_id": item["id"], "filename": f"[模板]{item['title']}",
            "chunk_index": 0, "total_chunks": 1,
            "upload_date": "2026-07-15 (built-in)", "char_count": len(item["content"]),
            "source": "builtin",
        })

    if new_docs:
        col.add(documents=new_docs, ids=new_ids, metadatas=new_metas)
        logger.info("Seeded %d trading templates (skipped %d existing)", len(new_docs), skipped)
    return {"seeded": len(new_docs), "skipped": skipped}


def seed_industry_kb() -> dict:
    """将行业分析模板写入 industry 知识库（幂等）。"""
    _ensure_init()
    col = _get_kb_collection("industry")
    existing = col.get()
    existing_ids = set(existing.get("ids", []))

    new_docs, new_ids, new_metas = [], [], []
    skipped = 0
    for item in _INDUSTRY_TEMPLATES:
        if item["id"] in existing_ids:
            skipped += 1
            continue
        new_ids.append(item["id"])
        new_docs.append(f"【{item['title']}】{item['content']}")
        new_metas.append({
            "doc_id": item["id"], "filename": f"[模板]{item['title']}",
            "chunk_index": 0, "total_chunks": 1,
            "upload_date": "2026-07-15 (built-in)", "char_count": len(item["content"]),
            "source": "builtin",
        })

    if new_docs:
        col.add(documents=new_docs, ids=new_ids, metadatas=new_metas)
        logger.info("Seeded %d industry templates (skipped %d existing)", len(new_docs), skipped)
    return {"seeded": len(new_docs), "skipped": skipped}


# ============================================================
#  公开 API（全部 kb_name 参数化）
# ============================================================

def list_kbs() -> list[dict]:
    """列出所有可用知识库及其统计信息。"""
    _ensure_init()
    try:
        reg = _client.get_collection(_REGISTRY_COLLECTION)
        data = reg.get(include=["metadatas"])
    except Exception:
        # registry 不存在时用内置定义回退
        data = {"ids": list(_BUILTIN_KBS.keys()),
                "metadatas": [{"display_name": v["display_name"],
                               "description": v["description"],
                               "builtin": True}
                              for v in _BUILTIN_KBS.values()]}

    result = []
    for i, kb_name in enumerate(data.get("ids", [])):
        meta = data["metadatas"][i] if data.get("metadatas") else {}
        # 统计该 knowledge base 的文档和切片数
        doc_ids, chunks = set(), 0
        try:
            col = _client.get_collection(f"kb_{kb_name}")
            col_data = col.get(include=["metadatas"])
            for m in col_data.get("metadatas", []):
                did = m.get("doc_id", "")
                if did:
                    doc_ids.add(did)
                chunks += 1
        except Exception:
            pass  # collection 尚不存在

        result.append({
            "name": kb_name,
            "display_name": meta.get("display_name", kb_name),
            "description": meta.get("description", ""),
            "builtin": meta.get("builtin", False),
            "doc_count": len(doc_ids),
            "chunk_count": chunks,
        })

    return result


def search_kb(query: str, kb_name: str = "accounting", top_k: int = 5) -> list[dict]:
    """语义检索指定知识库。

    Args:
        query: 自然语言查询
        kb_name: 知识库名 (accounting/industry/trading)
        top_k: 返回条数

    Returns:
        [{"content": "...", "source": "文件名", "score": 0.85}, ...]
    """
    _ensure_init()

    try:
        col = _get_kb_collection(kb_name)
    except Exception as e:
        logger.exception("Failed to get collection for kb '%s': %s", kb_name, e)
        return []

    try:
        results = col.query(
            query_texts=[query],
            n_results=top_k,
            include=["documents", "metadatas", "distances"],
        )
    except Exception as e:
        logger.exception("KB search failed for kb='%s' query='%s'", kb_name, query[:100])
        return []

    items = []
    if results.get("ids") and results["ids"][0]:
        for i in range(len(results["ids"][0])):
            dist = results.get("distances", [[1]])[0][i] if results.get("distances") else 1.0
            meta = results["metadatas"][0][i] if results.get("metadatas") else {}
            items.append({
                "content": results["documents"][0][i] if results.get("documents") else "",
                "source": meta.get("filename", "unknown"),
                "chunk": f"{meta.get('chunk_index', 0) + 1}/{meta.get('total_chunks', '?')}",
                "score": round(max(0, 1 - dist), 3),
            })

    return items


def upload_document(file_bytes: bytes, filename: str, kb_name: str = "accounting") -> dict:
    """上传并索引一份文档到指定知识库。

    Args:
        file_bytes: 文件原始字节
        filename: 原始文件名（用于识别格式）
        kb_name: 目标知识库名

    Returns:
        {"doc_id": "...", "filename": "...", "chunks": N, "size": bytes, "kb": kb_name}
    """
    _ensure_init()

    doc_id = str(uuid.uuid4())[:8]
    safe_name = f"{doc_id}_{filename}"
    file_path = os.path.join(_UPLOAD_DIR, safe_name)

    with open(file_path, "wb") as f:
        f.write(file_bytes)

    try:
        text = _parse_file(file_path, filename)
    except Exception:
        try:
            os.remove(file_path)
        except OSError:
            pass
        raise

    chunks = _chunk_text(text)
    if not chunks:
        try:
            os.remove(file_path)
        except OSError:
            pass
        raise ValueError(f"文档切分后无有效内容（文本{len(text)}字）: {filename}")

    col = _get_kb_collection(kb_name)
    chunk_ids = [f"{doc_id}_chunk_{i}" for i in range(len(chunks))]
    metadatas = [
        {
            "doc_id": doc_id,
            "filename": filename,
            "chunk_index": i,
            "total_chunks": len(chunks),
            "upload_date": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "char_count": len(chunks[i]),
            "kb": kb_name,
        }
        for i in range(len(chunks))
    ]

    try:
        col.add(documents=chunks, ids=chunk_ids, metadatas=metadatas)
    except Exception as e:
        try:
            col.delete(ids=chunk_ids)
        except Exception:
            pass
        try:
            os.remove(file_path)
        except OSError:
            pass
        raise RuntimeError(f"向量索引写入失败: {e}") from e

    logger.info("Document indexed: doc_id=%s, filename=%s, kb=%s, chunks=%d",
                doc_id, filename, kb_name, len(chunks))

    return {
        "doc_id": doc_id,
        "filename": filename,
        "kb": kb_name,
        "chunks": len(chunks),
        "size": len(file_bytes),
        "preview": text[:200] + ("..." if len(text) > 200 else ""),
    }


def list_documents(kb_name: str = "accounting") -> list[dict]:
    """列出指定知识库中所有已索引的文档（去重聚合）。"""
    _ensure_init()

    try:
        col = _get_kb_collection(kb_name)
        all_data = col.get(include=["metadatas"])
    except Exception as e:
        logger.exception("Failed to list documents in kb '%s': %s", kb_name, e)
        return []

    if not all_data.get("metadatas"):
        return []

    doc_map = {}
    for i, meta in enumerate(all_data["metadatas"]):
        did = meta.get("doc_id", "unknown")
        if did not in doc_map:
            doc_map[did] = {
                "doc_id": did,
                "filename": meta.get("filename", "unknown"),
                "upload_date": meta.get("upload_date", "unknown"),
                "total_chunks": 0,
                "total_chars": 0,
                "source": meta.get("source", "user"),
            }
        doc_map[did]["total_chunks"] += 1
        doc_map[did]["total_chars"] += meta.get("char_count", 0)

    return sorted(doc_map.values(), key=lambda d: d.get("upload_date", ""), reverse=True)


def delete_document(doc_id: str, kb_name: str = "accounting") -> bool:
    """删除指定知识库中的一份文档及其全部切片。

    Returns:
        True 如果成功删除，False 如果文档不存在
    """
    _ensure_init()

    try:
        col = _get_kb_collection(kb_name)
        all_data = col.get(include=["metadatas"])
    except Exception as e:
        logger.exception("Failed to query documents for deletion: kb=%s err=%s", kb_name, e)
        return False

    chunk_ids = []
    target_filename = None
    for i, meta in enumerate(all_data.get("metadatas", [])):
        if meta.get("doc_id") == doc_id:
            chunk_ids.append(all_data["ids"][i])
            target_filename = meta.get("filename")

    if not chunk_ids:
        logger.warning("Document not found for deletion: doc_id=%s kb=%s", doc_id, kb_name)
        return False

    try:
        col.delete(ids=chunk_ids)
    except Exception as e:
        logger.exception("Failed to delete chunks from ChromaDB: %s", e)
        return False

    if target_filename:
        safe_name = f"{doc_id}_{target_filename}"
        file_path = os.path.join(_UPLOAD_DIR, safe_name)
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except OSError as e:
            logger.warning("Failed to remove uploaded file %s: %s", file_path, e)

    logger.info("Document deleted: doc_id=%s, kb=%s, chunks=%d", doc_id, kb_name, len(chunk_ids))
    return True


def get_kb_stats(kb_name: str = "accounting") -> dict:
    """返回指定知识库统计信息。"""
    _ensure_init()
    docs = list_documents(kb_name)
    return {
        "kb_name": kb_name,
        "total_documents": len(docs),
        "total_chunks": sum(d["total_chunks"] for d in docs),
        "total_chars": sum(d["total_chars"] for d in docs),
        "upload_dir": _UPLOAD_DIR,
    }


# ---- 兼容旧API（与 rag.py 并存，不强制迁移） ----

def search_accounting(query: str, top_k: int = 5) -> list[dict]:
    """[兼容] 搜索会计准则知识库。与旧 rag.search_youzi 签名一致。"""
    return search_kb(query, "accounting", top_k)
